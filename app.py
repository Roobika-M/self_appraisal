from flask import Flask, request, render_template, redirect, url_for, make_response,flash,send_file
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
import logging
import os
import pandas as pd
from docx import Document
from docx2pdf import convert
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx.shared import Pt
import pythoncom
import win32com.client
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

app = Flask(__name__)
app.secret_key = "your_secret_key"

# app.config['SQLALCHEMY_DATABASE_URI'] = 'postgresql://postgres:harshu7564@localhost/KITE_STAFF'
# app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# db = SQLAlchemy(app)

# class userlo(db.Model):
#     id = db.Column(db.Integer, primary_key=True)
#     usernam = db.Column(db.String(70), nullable=False, unique=True)
#     password = db.Column(db.String(120), nullable=False)

#     def __repr__(self):
#         return f"<User {self.usernam}>"

# class staff(db.Model):
#     id = db.Column(db.Integer, primary_key=True)
#     name = db.Column(db.String(100), nullable=False)
#     position = db.Column(db.String(50))

#     def __repr__(self):
#         return f"<Staff {self.name}>"

# Render the login page
@app.route('/')
def home():
    return render_template('login.html', error=None)

# Handle form submission
@app.route('/login', methods=['POST', 'GET'])
def login():
    if request.method == 'POST':  
        username = request.form.get('username')
        password = request.form.get('password')

        if not username or not password:
            error = "Please enter both username and password."
            return render_template('login.html', error=error)

        # usr = userlo.query.filter_by(usernam=username).first()
        usr=username

        # if usr and password == usr.password:
        if usr and password == "admin":
            return redirect(url_for('upload'))
        else:
            error = "Please enter correct username and password."
            return render_template('login.html', error=error)

    return render_template('login.html')
staffname = ''
detaillist = []
excel_path = ''

@app.route('/upload', methods=['POST', 'GET'])
def upload():
    global staffname, detaillist, excel_path

    if request.method == 'POST':
        print("Form submitted!")  # Debugging
        
        # Get form data
        name = request.form.get('name')
        designation = request.form.get('designation')
        department = request.form.get('dept')
        emp_id = request.form.get('empid')

        # Validate form data
        if not all([name, designation, department, emp_id]):
            print("Missing form data!")  # Debugging
            return render_template('upload.html', error="Please fill in all details.")
        staffname = name
        detaillist = [name, designation, department, emp_id]

        # Get uploaded files
        excel_file = request.files.get('excel_file')
        template_file = request.files.get('template_file')
        
        if not excel_file or excel_file.filename == '' or not template_file or template_file.filename == '':
            print("Missing files!")  # Debugging
            return render_template('upload.html', error="Please upload both Excel and Template files.")

        # Save the files
        upload_folder = os.getcwd()
        excel_path = os.path.join(upload_folder, excel_file.filename)
        template_path = os.path.join(upload_folder, template_file.filename)
        
        excel_file.save(excel_path)
        template_file.save(template_path)

        print(f"Files saved: Excel - {excel_path}, Template - {template_path}")  # Debugging

        # try:
        processing(excel_path, staffname, template_path)  # Pass template_path to processing function
        # except Exception as e:
            # print(f"Error processing files: {e}")  # Debugging
            # return render_template('upload.html', error="Error processing files.")

        print("Redirecting to download page...")  # Debugging
        return redirect(url_for("download_path"))

    return render_template('upload.html')
    

@app.route('/data', methods=['POST', 'GET'])
def data():
    if request.method == 'POST':
        id = request.form.get('id')
        username = request.form.get('username')
        name = request.form.get('name')
        password = request.form.get('password')
        position = request.form.get('posi')

        # if not all([username, name, password, position]):
        #     error = "Please enter all details."
        #     return render_template('data.html', error=error)
        # existing_user = userlo.query.filter_by(usernam=username).first()

        # if existing_user:
        #     error = "Username exist."
        #     return render_template('data.html', error=error)
        
        # try:
        #     new_user = userlo(id=id,usernam=username, password=password)
        #     new_staff = staff(id=id,name=name, position=position)

        #     db.session.add(new_user)
        #     db.session.add(new_staff)
        #     db.session.commit()

        #     flash("User added successfully!", "success")
        #     return redirect(url_for('login'))
        # except Exception as e:
        #     db.session.rollback()
        #     flash(f"Error: {str(e)}", "error")
        #     return redirect(url_for('data'))

    return render_template('data.html')

'''logging.basicConfig()
logging.getLogger('sqlalchemy.engine').setLevel(logging.INFO)'''

@app.route('/download/<file_type>', methods=['GET'])
def download(file_type):
    upload_folder = os.getcwd()  # Get current working directory
    docx_filename = os.path.join(upload_folder, "filled_template.docx")
    pdf_filename = os.path.join(upload_folder, "filled_template.pdf")
    docx_corrective_filename = os.path.join(upload_folder, "appfilled_template.docx")
    pdf_corrective_filename = os.path.join(upload_folder, "appfilled_template.pdf")

    if file_type == "docx":
        if not os.path.exists(docx_filename):
            return "File not found", 404
        return send_file(docx_filename, as_attachment=True)

    elif file_type == "pdf":
        convert_docx_to_pdf(docx_filename, pdf_filename)
        if not os.path.exists(pdf_filename):
            return "PDF conversion failed", 500
        return send_file(pdf_filename, as_attachment=True)

    elif file_type == "docx_corrective":
        if not os.path.exists(docx_corrective_filename):
            return "File not found", 404
        return send_file(docx_corrective_filename, as_attachment=True)

    elif file_type == "pdf_corrective":
        convert_docx_to_pdf(docx_corrective_filename, pdf_corrective_filename)
        if not os.path.exists(pdf_corrective_filename):
            return "PDF conversion failed", 500
        return send_file(pdf_corrective_filename, as_attachment=True)

    return "Invalid file type", 400

@app.route('/download_path')
def download_path():
    global staffname, research, selfm, mentor, academics, hod
    print(type(research), type(selfm), type(mentor), type(academics), type(hod))
    total_score = int(research) + int(selfm) + int(mentor) + int(academics) + int(hod)    
    return render_template(
        "download.html", 
        name=staffname, 
        research=research, 
        selfm=selfm,
        mentor=mentor, 
        academics=academics,  # Add academics score
        hod=hod,
        total_score=total_score
    )


def convert_docx_to_pdf(docx_path, pdf_path):
    pythoncom.CoInitialize()  # Fixes COM errorWord
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False  # Run in background

    try:
        docx_path = os.path.abspath(docx_path)  # Ensure absolute path
        pdf_path = os.path.abspath(pdf_path)  # Ensure absolute path
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # Convert DOCX to PDF
        doc.Close()
    except Exception as e:
        print(f"Error: {e}")
    finally:
        word.Quit()  
        pythoncom.CoUninitialize()  


def find_header_row(excel_path, sheet_name):
    # Read the first few rows to find the header
    df_temp = pd.read_excel(excel_path, sheet_name=sheet_name, nrows=15)
    
    # Look for "Faculty Name" or "Name of the faculty" in each row
    for idx, row in df_temp.iterrows():
        row_values = [str(val).lower().strip() for val in row if pd.notna(val)]
        if 'faculty name' in row_values or 'name of the faculty' in row_values:
            return idx
    
    return None 


#################################### Load the Excel file
# Load the Word document
def processing(excel_path, staffname, template_path):  # Add template_path parameter
    doc = Document("template.docx")
    doc1 = Document(template_path)  # Use the uploaded template file
    name = staffname
    global m, n, research,selfm,mentor,academics,hod  
    m = 0
    n = 0
    research = 0
    academics = 0
    selfm = 0
    mentor = 0
    hod = 0
    sheet_names = pd.ExcelFile(excel_path).sheet_names
    ###############################################################
    scores=[]

    # Access the second tables (Academics)   
    source_table = doc1.tables[1]
    destination_table = doc.tables[1]

    # Copy from source row 2 onward (skip first 2 header rows)
    scores = [0, 0, 0, 0, 0, 0]

    # Helper: Convert text to float
    def get_float(cell_text):
        try:
            return float(cell_text.strip())
        except:
            return 0

    # Helper: Get grade for each score based on rules
    def get_grade_1(value):
        if value > 95:
            return 5
        elif 90 <= value <= 95:
            return 4
        elif 80 <= value < 90:
            return 3
        elif 70 <= value < 80:
            return 2
        elif 60 <= value < 70:
            return 1
        elif 50 <= value < 60:
            return 0
        else:
            return -1

    def get_grade_2(value):
        if 0 < value <= 2:
            return 1
        elif 3 <= value <= 4:
            return 2
        elif 5 <= value <= 6:
            return 3
        elif 7 <= value <= 9:
            return 4
        else:
            return 5

    def get_grade_negative(value):
        if 0 < value <= 10:
            return -1
        elif 11 <= value <= 20:
            return -2
        elif 21 <= value <= 30:
            return -3
        elif 31 <= value <= 40:
            return -4
        else:
            return -5

    # Loop through rows in source table starting from row 2
    for i in range(2, len(source_table.rows)):
        row = source_table.rows[i]
        first_cell = row.cells[0].text.strip()

        # If it's the "Total/Average" row, handle specially
        if first_cell.lower() == "total/average":
            nos = i - 3  # number of valid data rows
            print(nos)

            for j in range(i, i + 2):
                print(scores)
                
                source_row = source_table.rows[j]
                if j >= len(destination_table.rows):
                    destination_table.add_row()
                new_row = destination_table.rows[j]

                # Ensure enough cells
                while len(new_row.cells) < len(source_row.cells):
                    new_row._tr.add_tc()

                # Merge first 4 columns
                merged_cell = new_row.cells[0].merge(new_row.cells[1])
                merged_cell = merged_cell.merge(new_row.cells[2])
                merged_cell = merged_cell.merge(new_row.cells[3])

                if j == i:
                    new_row.cells[3].text = "Total/Average"
                    avg = (scores[0] // nos) if nos > 0 else 0
                    new_row.cells[4].text = f"{avg:.2f}"
                    for k in range(1,6):
                        new_row.cells[k + 4].text = f"{scores[k]:.2f}"
                elif j == i + 1:
                    new_row.cells[3].text = "Marks(Ref guideline for awarding score)"
                    new_row.cells[4].text = str(get_grade_1(scores[0] / nos))
                    new_row.cells[5].text = str(get_grade_2(scores[1]))
                    new_row.cells[6].text = str(get_grade_2(scores[2]))
                    new_row.cells[7].text = str(get_grade_2(scores[3]))
                    new_row.cells[8].text = str(get_grade_negative(scores[4]))
                    new_row.cells[9].text = str(get_grade_negative(scores[5]))
            break  # Done after adding Total/Average and Marks rows

        # Normal row copy
        if i >= len(destination_table.rows):
            destination_table.add_row()

        new_row = destination_table.rows[i]
        while len(new_row.cells) < len(row.cells):
            new_row._tr.add_tc()

        # Copy and accumulate scores
        for j in range(len(row.cells)):
            text = row.cells[j].text.strip()
            new_row.cells[j].text = text
            if j >= 4 and j <= 9:
                scores[j - 4] += get_float(text)
            
    def get_total_academics_score(scores, nos):
        total = 0
        if nos > 0:
            total += get_grade_1(scores[0] / nos)  # First column score
        total += get_grade_2(scores[1])  # Second column score
        total += get_grade_2(scores[2])  # Third column score 
        total += get_grade_2(scores[3])  # Fourth column score
        total += get_grade_negative(scores[4])  # Fifth column score
        total += get_grade_negative(scores[5])  # Sixth column score
        return total

    # After calculating the marks row:
    academics = get_total_academics_score(scores, nos)  # Store the total academics score
    u1,academics = str(academics), academics

######################################################################################################
    global r1_1, r2_1, r3_1, r4_1, r5_1, r6_1, r7_1, r8_1, r9_1, r10_1, r11_1, r12_1, r13_1
    r1_1, r2_1, r3_1, r4_1, r5_1, r6_1, r7_1, r8_1, r9_1, r10_1, r11_1, r12_1, r13_1 = 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0    
    if "Journal Publication" in sheet_names:
        print(find_header_row(excel_path, "Journal Publication"))
        df_journal = pd.read_excel(excel_path, sheet_name="Journal Publication", skiprows=find_header_row(excel_path, "Journal Publication")+1)

        # Fix column names (remove spaces)
        df_journal.columns = df_journal.columns.str.strip()
        df_filtered=[]
        df_journal["Faculty Name"] = df_journal["Faculty Name"].ffill()
        df_filtered = df_journal[df_journal["Faculty Name"].str.strip() == name]
        m = 0
        if not df_filtered.empty:
            table3_index = 3  
            table3 = doc.tables[table3_index]

            start_row = 1

            n=0
            
            # Fill the table3 with Excel data
            ###################table 1-journal###########################
            for i, (_, row) in enumerate(df_filtered.iterrows()):
                row_index = start_row+i
                if i+2 >= len(table3.rows):  
                    table3.add_row()  # Add rows if needed

                table3.cell(row_index+1, 0).text = str(i+1)  # Serial No.
                table3.cell(row_index+1, 1).text = str(row.get("Paper Title", "-"))
                table3.cell(row_index+1, 2).text = str(row.get("Journal Name", "-"))
                table3.cell(row_index+1, 3).text = str(row.get("Year of Publication", "-"))
                table3.cell(row_index+1, 4).text = str(row.get("ISSN", "-"))
                table3.cell(row_index+1, 5).text = str(row.get("Web Link", "-"))
                table3.cell(row_index+1, 6).text = str(row.get("Impact Factor", "-"))
                if row.get("Impact Factor", "-")!='-':
                    if row.get("Impact Factor", "-")>3:
                        n+=3
                        r2_1+=3
                    elif row.get("Impact Factor", "-")>1.5 and row.get("Impact Factor", "-")<=3:
                        n+=2
                        r3_1+=2
                    elif row.get("Impact Factor", "-")>=1 and row.get("Impact Factor", "-")<=1.5:
                        n+=1
                        r4_1+=1
                n+=2 
                r1_1+2
                

            
            row_index=-1
            table3.add_row()
            table3.rows[row_index].cells[0].merge(table3.rows[row_index].cells[-2])
            paragraph = table3.cell(-1, 5).paragraphs[0]
            run = paragraph.add_run("Total score")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            table3.cell(-1, 6).text = str(n)
            m = m+n
    #######################table2-book au#################################  
    if "Book Publication" in sheet_names:

        n = 0
        df_bookpub = pd.read_excel(excel_path, sheet_name="Book Publication", skiprows=find_header_row(excel_path, "Book Publication")+1)

        df_bookpub.columns = df_bookpub.columns.str.strip()
        df_filtered=[]
        df_bookpub["Faculty Name"] = df_bookpub["Faculty Name"].ffill()
        df_filtered = df_bookpub[df_bookpub["Faculty Name"].str.strip() == name]
        if not df_filtered.empty:
            table4_index = 4  
            table4 = doc.tables[table4_index]
            start_row = 1

            for i, (_, row) in enumerate(df_filtered.iterrows()):
                row_index = start_row+i
                if i+2 >= len(table4.rows):  
                    table4.add_row()  # Add rows if needed

                table4.cell(row_index+1, 0).text = str(i+1)  # Serial No.
                table4.cell(row_index+1, 1).text = str(row.get("Book Title", "-"))
                table4.cell(row_index+1, 2).text = str(row.get("Publication Name", "-"))
                table4.cell(row_index+1, 3).text = str(row.get("Date of Publication", "-"))  # Fixed Date Mapping
                table4.cell(row_index+1, 4).text = str(row.get("ISBN", "-"))
                table4.cell(row_index+1, 5).text = str(row.get("Description", "-"))  
        ################################table3-invited lectures##########################################
        # no data shiiiiiiiiiiiiiiiiiiiiiiii
        #######################################table4,5-conference##################################
    if "Conferences" in sheet_names:
        n = 0
        df_conference = pd.read_excel(excel_path, sheet_name="Conferences", skiprows=find_header_row(excel_path, "Conferences")+1)
        df_filtered=[]
        df_conference.columns = df_conference.columns.str.strip()
        df_conference["Faculty Name"] = df_conference["Faculty Name"].ffill()

        df_filtered = df_conference[df_conference["Faculty Name"].str.strip() == name]
        if not df_filtered.empty:
        # Get table references
            table6_index = 6  # International conference table
            table7_index = 7  # National conference table
            table6 = doc.tables[table6_index]
            table7 = doc.tables[table7_index]

            start_row = 1
            table = 0
            for i, (_, row) in enumerate(df_filtered.iterrows()):
                conference_type = str(row["Conference Type"]).strip().lower()
                if conference_type == "international":
                    table = table6
                    organized_by = row.get("Organized By", "-")
                    n += 2
                    r8_1+=2
                else:
                    table = table7
                    organized_by = row.get("Organized By", "-")
                    n += 1
                    r9+=1
                row_index = start_row + i
                
                if i + 2 >= len(table.rows):
                    table.add_row()  # Add rows if needed
                
                table.cell(row_index + 1, 0).text = str(i + 1)  # Serial No.
                table.cell(row_index + 1, 1).text = str(row.get("Paper Title", "-"))
                table.cell(row_index + 1, 2).text = organized_by
                table.cell(row_index + 1, 3).text = str(row.get("From Date", "-"))
                table.cell(row_index + 1, 4).text = str(row.get("Place", "-"))
                table.cell(row_index + 1, 5).text = str(row.get("Role", "-"))
            if table!=0:
                row_index=-1
                table.add_row()
                table.rows[row_index].cells[0].merge(table.rows[row_index].cells[-2])
                paragraph = table.cell(-1, 4).paragraphs[0]
                run = paragraph.add_run("Total score")
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                table.cell(-1, 5).text = str(n)
            m = m+n
    ##################################table6 research grant##############################################
    if "Research Grant" in sheet_names:
        n = 0
        total_amt=0
        df_filtered=[]
        df_research = pd.read_excel(excel_path, sheet_name="Research Grant", skiprows=find_header_row(excel_path, "Research Grant")+1)

        df_research.columns = df_research.columns.str.strip()

        df_research["Faculty Name"] = df_research["Faculty Name"].ffill()
        df_filtered = df_research[df_research["Faculty Name"].str.strip() == name]
        if not df_filtered.empty:
            table7_index = 7  
            table7 = doc.tables[table7_index]
            start_row = 1

            for i, (_, row) in enumerate(df_filtered.iterrows()):
                row_index = start_row+i
                if i+2 >= len(table7.rows):  
                    table7.add_row()  # Add rows if needed
                if str(row.get("Coordinator", "-"))=="applied":
                    table7.cell(row_index+1, 0).text = str(i+1)  # Serial No.
                    table7.cell(row_index+1, 1).text = str(row.get("Coordinator", "-"))
                    table7.cell(row_index+1, 2).text = str(row.get("Title", "-")) 
                    table7.cell(row_index+1, 3).text = str(row.get("Type", "-"))
                    table7.cell(row_index+1, 4).text = str(row.get("Funding Agent", "-"))  
                    table7.cell(row_index+1, 5).text = str(row.get("Amount", "-"))  
                    table7.cell(row_index+1, 6).text = str(row.get("Applied On", "-"))  

                    if row.get("Amount", "-")!="-":
                        total_amt+=row.get("Amount", "-")
            if row.get("Amount", "-")>1000000:
                n+=(row.get("Amount", "-")//1000000)*2
                r10+=n
            row_index=-1
            table7.add_row()
            table7.rows[row_index].cells[0].merge(table7.rows[row_index].cells[-2])
            paragraph = table7.cell(-1, 5).paragraphs[0]
            run = paragraph.add_run("Total score")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            table7.cell(-1, 6).text = str(n)
            m = m+n
    #########################################table 7-seminar#########################################
    if "Research Grant" in sheet_names:
        n = 0
        df_seminar = pd.read_excel(excel_path, sheet_name="Research Grant", skiprows=find_header_row(excel_path, "Research Grant")+1)
        df_filtered=[]
        df_seminar.columns = df_seminar.columns.str.strip()

        df_seminar["Faculty Name"] = df_seminar["Faculty Name"].ffill()
        df_filtered = df_seminar[df_seminar["Faculty Name"].str.strip() == name]
        if not df_filtered.empty:
            table9_index = 9 
            table9 = doc.tables[table9_index]
            start_row = 1
            if not df_filtered.empty:
                for i, (_, row) in enumerate(df_filtered.iterrows()):
                    row_index = start_row+i
                    if i+2 >= len(table9.rows):  
                        table9.add_row()  # Add rows if needed
                    if str(row.get("Coordinator", "-"))!="applied":
                        table9.cell(row_index+1, 0).text = str(i+1)  # Serial No.
                        table9.cell(row_index+1, 1).text = str(row.get("Coordinator", "-"))
                        table9.cell(row_index+1, 2).text = str(row.get("Title", "-")) 
                        table9.cell(row_index+1, 3).text = str(row.get("Type", "-"))
                        table9.cell(row_index+1, 4).text = str(row.get("Funding Agent", "-"))  
                        table9.cell(row_index+1, 5).text = str(row.get("Amount", "-"))  
                        table9.cell(row_index+1, 6).text = str(row.get("Applied On", "-")) 

                        if row.get("Amount", "-")>50000:
                            n+=(row.get("Amount", "-")//50000)
                            r11_1+=n###dout 

                        # if row.get("Amount", "-")!="-":
                        #     total_amt+=row.get("Amount", "-")
                
                row_index=-1
                table9.add_row()
                table9.rows[row_index].cells[0].merge(table9.rows[row_index].cells[-2])
                paragraph = table9.cell(-1, 5).paragraphs[0]
                run = paragraph.add_run("Total score")
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                table9.cell(-1, 6).text = str(n)
                m = m+n
    ###############################################table 7,9-patent##################
    if "Patents" in sheet_names:
        n = 0
        df_patent = pd.read_excel(excel_path, sheet_name="Patents")
        df_filtered=[]
        df_patent.columns = df_patent.columns.str.strip()

        df_patent["Faculty name"] = df_patent["Faculty name"].ffill()
        df_filtered = df_patent[df_patent["Faculty name"].str.strip() == name]
        if not df_filtered.empty:
            table10_index = 10
            table10 = doc.tables[table10_index]

            start_row = 1

            from docx.shared import Pt

            def clear_and_write(cell, text):
                """ Clears the cell and writes new text with formatting. """
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.clear()  # Ensure old text is removed
                cell.text = ""  # Extra clearing step
                run = cell.paragraphs[0].add_run(text)
                run.font.size = Pt(11)  # Ensure text is visible
                run.font.bold = True  # Make text bold to check visibility

            for i, (_, row) in enumerate(df_filtered.iterrows()):
                row_index = start_row + i
                if row_index + 1 >= len(table10.rows):  
                    table10.add_row()  # Add new rows if needed

                # Debugging: Print before writing
                print(f"Writing to row {row_index+1}: {row.to_dict()}")

                # Fill Serial No.
                clear_and_write(table10.cell(row_index + 1, 0), str(i + 1))

                # Fill Patent Title
                title = str(row.get("Title", "-")).strip()
                clear_and_write(table10.cell(row_index + 1, 1), title)

                # Determine Filing/Publishing Date
                status = str(row.get("Status", "")).strip().lower()
                date_value = str(row.get("Date", "-"))

                if status == "filed":
                    clear_and_write(table10.cell(row_index + 1, 2), date_value)  # Date of Filing
                    clear_and_write(table10.cell(row_index + 1, 3), "-")  # No publishing date
                elif status == "published":
                    clear_and_write(table10.cell(row_index + 1, 2), "-")  # No filing date
                    clear_and_write(table10.cell(row_index + 1, 3), date_value)  # Date of Publish
                    n+=5
                    r12_1+=5
                else:
                    clear_and_write(table10.cell(row_index + 1, 2), "-")
                    clear_and_write(table10.cell(row_index + 1, 3), "-")

                # Fill Other Details
                clear_and_write(table10.cell(row_index + 1, 4), str(row.get("Status", "-")))

            row_index=-1
            table10.add_row()
            table10.rows[row_index].cells[0].merge(table10.rows[row_index].cells[-2])
            paragraph = table10.cell(-1, 3).paragraphs[0]
            run = paragraph.add_run("Total score")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            table10.cell(-1, 4).text = str(n)
            m = m+n
    research = m

    #####################################table 10-consultancy####################
    #table consultancy no data da
    ########################################table 11-workshop#################
    global p1_1,p2_1,p3_1,p4_1,p5_1,p6_1,p7_1
    p1_1,p2_1,p3_1,p4_1,p5_1,p6_1,p7_1=0,0,0,0,0,0,0
    m = 0
    n = 0
    if "Workshop" in sheet_names:
        df_workshop = pd.read_excel(excel_path, sheet_name="Workshop", skiprows=find_header_row(excel_path, "Workshop")+1)

        df_workshop.columns = df_workshop.columns.str.strip()
        df_filtered=[]
        df_workshop["Faculty Name"] = df_workshop["Faculty Name"].ffill()
        df_filtered = df_workshop[df_workshop["Faculty Name"].str.strip() == name]
        if not df_filtered.empty:
            table13_index = 14
            table13 = doc.tables[table13_index]
            start_row = 1

            for i, (_, row) in enumerate(df_filtered.iterrows()):
                
                from_date = str(row.get("From Date","-"))
                to_date = str(row.get("To Date","-"))
                row_index = start_row+i
                if i+2 >= len(table13.rows):  
                    table13.add_row()  # Add rows if needed
                if str(row.get("Role")) == "attended":
                    table13.cell(row_index+1, 0).text = str(i+1)  # Serial No.
                    table13.cell(row_index+1, 1).text = str(row.get("Topic", "-"))
                    table13.cell(row_index+1, 2).text = f"{from_date} to {to_date}"
                    table13.cell(row_index+1, 3).text = str(row.get("Description", "-"))  
                    table13.cell(row_index+1, 4).text = str(row.get("Venue", "-")) 
                    if n<3:
                        n+=1
                        p1_1+=1

            row_index=-1
            table13.add_row()
            table13.rows[row_index].cells[0].merge(table13.rows[row_index].cells[-2])
            paragraph = table13.cell(-1, 3).paragraphs[0]
            run = paragraph.add_run("Total score")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            table13.cell(-1, 4).text = str(n)
            
            m = m+n

    ###################################table 12-skill dev####################################
    if "Faculty Internship" in sheet_names:
        n = 0
        df_filtered=[]
        df_develop = pd.read_excel(excel_path, sheet_name="Faculty Internship", skiprows=find_header_row(excel_path, "Faculty Internship")+1)
        df_develop.columns = df_develop.columns.str.strip()
        possible_names=["Faculty Name","Faculty Name"] 
        selectedcol=next((col for col in possible_names if col in df_develop.columns),None)
        if selectedcol:
            df_develop[selectedcol] = df_develop[selectedcol].ffill()
            df_filtered = df_develop[df_develop[selectedcol].str.strip() == name]
        if not df_filtered.empty:
            table14_index = 15
            table14 = doc.tables[table14_index]
            start_row = 1

            for i, (_, row) in enumerate(df_filtered.iterrows()):
                from_date = str(row.get("From Date","-"))
                to_date = str(row.get("To Date","-"))
                row_index = start_row+i
                if i+2 >= len(table14.rows):  
                    table14.add_row()  # Add rows if needed

                table14.cell(row_index+1, 0).text = str(i+1)  # Serial No.
                table14.cell(row_index+1, 1).text = str(row.get("FDP Name", "-"))
                table14.cell(row_index+1, 2).text = f"{from_date} to {to_date}"
                table14.cell(row_index+1, 3).text = str(row.get("Description", "-"))  
                table14.cell(row_index+1, 4).text = str(row.get("National or International", "-"))
                n += 3
                p2_1+=3

            row_index=-1
            table14.add_row()
            table14.rows[row_index].cells[0].merge(table14.rows[row_index].cells[-2])
            paragraph = table14.cell(-1, 3).paragraphs[0]
            run = paragraph.add_run("Total score")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            table14.cell(-1, 4).text = str(n)
            
            m = m+n
    #############################table 13-mooc&nptel#############################################
    if "MOOC Course" in sheet_names:
        n = 0
        df_filtered=[]
        df_mooc = pd.read_excel(excel_path,sheet_name="MOOC Course", skiprows=find_header_row(excel_path, "MOOC Course")+1)
        df_mooc.columns = df_mooc.columns.str.strip()
        possible_names=["Faculty Name","Faculty Name"] 
        selectedcol=next((col for col in possible_names if col in df_mooc.columns),None)
        if selectedcol:
            df_mooc[selectedcol] = df_mooc[selectedcol].ffill()
            df_filtered = df_mooc[df_mooc[selectedcol].str.strip() == name]
        if not df_filtered.empty: 
            table15_index = 16
            table15 = doc.tables[table15_index] 
            start_row = 1
            for i, (_, row) in enumerate(df_filtered.iterrows()):
                from_date = str(row.get("From Date","-"))
                to_date = str(row.get("To Date","-"))
                row_index = start_row+i
                if i+2 >= len(table15.rows):  
                    table15.add_row()  # Add rows if needed

                table15.cell(row_index+1, 0).text = str(i+1)  # Serial No.
                table15.cell(row_index+1, 1).text = str(row.get("Coure Title", "-"))
                table15.cell(row_index+1, 2).text = str(row.get("Course Type", "-"))
                table15.cell(row_index+1, 3).text = f"{from_date} to {to_date}"
                table15.cell(row_index+1, 4).text = str(row.get("Duration", "-")) 
                table15.cell(row_index+1, 5).text = str(row.get("Awards","-"))
                if n < 4:
                    n += 2
                    p3_1+=2

            row_index=-1
            table15.add_row()
            table15.rows[row_index].cells[0].merge(table15.rows[row_index].cells[-2])
            paragraph = table15.cell(-1, 4).paragraphs[0]
            run = paragraph.add_run("Total score")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            table15.cell(-1, 5).text = str(n)
            
            m = m+n
    ##################################table 14-mou##################################
    if "MoU" in sheet_names:
        n = 0
        df_filtered=[]
        df_mou = pd.read_excel(excel_path,sheet_name="MoU", skiprows=find_header_row(excel_path, "MoU")+1)
        df_mou.columns = df_mou.columns.str.strip()
        possible_names=["Faculty Name","Faculty Name"] 
        selectedcol=next((col for col in possible_names if col in df_mou.columns),None)
        if selectedcol:
            df_mou[selectedcol] = df_mou[selectedcol].ffill()
            df_filtered = df_mou[df_mou[selectedcol].str.strip() == name]
        if not df_filtered.empty:
            table16_index = 17
            table16 = doc.tables[table16_index]
            start_row = 1
            for i, (_, row) in enumerate(df_filtered.iterrows()):
                from_date = str(row.get("From Date","-"))
                to_date = str(row.get("To Date","-"))
                row_index = start_row+i
                if i+2 >= len(table16.rows):  
                    table16.add_row()  # Add rows if needed

                table16.cell(row_index+1, 0).text = str(i+1)  # Serial No.
                table16.cell(row_index+1, 1).text = str(row.get(selectedcol, "-"))
                table16.cell(row_index+1, 2).text = str(row.get("Company Name", "-"))
                table16.cell(row_index+1, 3).text = f"{from_date} to {to_date}"
                table16.cell(row_index+1, 4).text = str(row.get("Industry SPOC", "-")) 
                table16.cell(row_index+1, 5).text = str(row.get("Duration","-"))
                n += 1
                p4_1+=1

            row_index=-1
            table16.add_row()
            table16.rows[row_index].cells[0].merge(table14.rows[row_index].cells[-2])
            paragraph = table16.cell(-1, 4).paragraphs[0]
            run = paragraph.add_run("Total score")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            table16.cell(-1, 5).text = str(n)
            
            m = m+n
    ###################################table 16-spl contribute##################################

    # df_awards = pd.read_excel(excel_path,sheet_name="Extension Activities", skiprows=5)
    # df_awards.columns = df_awards.columns.str.strip()
    # df_awards["Faculty Name"] = df_awards["Faculty Name"].ffill()
    # df_filtered = df_awards[df_awards["Faculty Name"].str.strip() == name]
    # table17_index = 17
    # table17 = doc.tables[table17_index]
    # start_row = 1
    # for i, (_, row) in enumerate(df_filtered.iterrows()):
    #     from_date = str(row.get("From Date","-"))
    #     to_date = str(row.get("To Date","-"))
    #     row_index = start_row+i
    #     if i+2 >= len(table17.rows):  
    #         table17.add_row()  # Add rows if needed

    #     table17.cell(row_index+1, 0).text = str(i+1)  # Serial No.
    #     table17.cell(row_index+1, 1).text = str(row.get("Name of the Event", "-"))
    #     table17.cell(row_index+1, 2).text = f"{from_date} to {to_date}"
    #     table17.cell(row_index+1, 3).text = str(row.get("Recognition", "-"))  
    #     table17.cell(row_index+1, 4).text = str(row.get("Award", "-"))
    #     table17.cell(row_index+1, 5).text = str(row.get("Description","-"))
    ###################################table 16-no of conference,workshop,hack###############################################
    if "Workshops" in sheet_names:
        n = 0
        df_workshop = pd.read_excel(excel_path, sheet_name="Workshops", skiprows=find_header_row(excel_path, "Workshops")+1)
        df_filtered=[]
        df_workshop.columns = df_workshop.columns.str.strip()

        df_workshop["Faculty Name"] = df_workshop["Faculty Name"].ffill()
        df_filtered = df_workshop[(df_workshop["Faculty Name"].str.strip() == name) & (df_workshop["Role"].fillna("").str.strip() == "conducted")]
        if not df_filtered.empty:
            table17_index = 19
            table17 = doc.tables[table17_index]
            start_row = 1

            for i, (_, row) in enumerate(df_filtered.iterrows()):
                
                from_date = str(row.get("From Date","-"))
                to_date = str(row.get("To Date","-"))
                row_index = start_row+i
                if i+2 >= len(table17.rows):  
                    table17.add_row()  # Add rows if needed
                if str(row.get("Role", "-"))=="conducted":
                    table17.cell(row_index+1, 0).text = str(i+1)  # Serial No.
                    table17.cell(row_index+1, 1).text = str(row.get("Topic", "-"))
                    table17.cell(row_index+1, 2).text = str(row.get("Department", "-"))
                    table17.cell(row_index+1, 3).text = f"{from_date} to {to_date}"
                    table17.cell(row_index+1, 4).text = str(row.get("No of Students", "-"))  
                    table17.cell(row_index+1, 5).text = str(row.get("Venue", "-"))  
                    table17.cell(row_index+1, 6).text = str(row.get("Description", "-")) 

                    n += 0.5
                    p6_1+=0.5

            row_index=-1
            table17.add_row()
            table17.rows[row_index].cells[0].merge(table17.rows[row_index].cells[-2])
            paragraph = table17.cell(-1, 5).paragraphs[0]
            run = paragraph.add_run("Total score")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            table17.cell(-1, 6).text = str(n)
            
            m = m+n
    ################################################table 17-expert visit###############################
    if "Guest Lectures" in sheet_names:
        n = 0
        df_filtered=[]
        df_experts = pd.read_excel(excel_path, sheet_name="Guest Lectures", skiprows=find_header_row(excel_path, "Guest Lectures")+1)

        df_experts.columns = df_experts.columns.str.strip()

        df_experts["Faculty Name"] = df_experts["Faculty Name"].ffill()
        df_filtered = df_experts[df_experts["Faculty Name"].str.strip() == name]
        if not df_filtered.empty:
            table19_index = 20
            table19 = doc.tables[table19_index]
            start_row = 1

            for i, (_, row) in enumerate(df_filtered.iterrows()):
                
                from_date = str(row.get("From Date","-"))
                to_date = str(row.get("To Date","-"))
                row_index = start_row+i
                if i+2 >= len(table19.rows):  
                    table19.add_row()  # Add rows if needed

                table19.cell(row_index+1, 0).text = str(i+1)  # Serial No.
                table19.cell(row_index+1, 1).text = str(row.get("Chief Guest Name", "-"))
                table19.cell(row_index+1, 2).text = str(row.get("Address", "-"))
                table19.cell(row_index+1, 3).text = str(row.get("Topic Name","-"))
                table19.cell(row_index+1, 4).text = f"{from_date} to {to_date}"  
                table19.cell(row_index+1, 5).text = str(row.get("Description", "-"))  
                table19.cell(row_index+1, 6).text = str(row.get("Topic Delivered", "-")) 
                n += 1
                p7_1+=1


            row_index=-1
            table19.add_row()
            table19.rows[row_index].cells[0].merge(table19.rows[row_index].cells[-2])
            paragraph = table19.cell(-1, 5).paragraphs[0]
            run = paragraph.add_run("Total score")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            table19.cell(-1, 6).text = str(n)
            
            m = m+n
    selfm = m
    #####################table 17-###########################################################
    m = 0
    n=0
    global s1_1,s2_1,s3_1,s4_1,s5_1
    s1_1,s2_1,s3_1,s4_1,s5_1=0,0,0,0,0
    if "Project Guided or Mentoring" in sheet_names:
        df_filtered=[]
        df_project = pd.read_excel(excel_path, sheet_name="Project Guided or Mentoring",skiprows=find_header_row(excel_path, "Project Guided or Mentoring")+1)

        df_project.columns = df_project.columns.str.strip()

        df_project["Faculty Name"] = df_project["Faculty Name"].ffill()
        df_filtered = df_project[df_project["Faculty Name"].str.strip() == name]
        if not df_filtered.empty:
            table21_index = 22
            table21 = doc.tables[table21_index]
            start_row = 1

            for i, (_, row) in enumerate(df_filtered.iterrows()):
                
                from_date = str(row.get("From Date","-"))
                to_date = str(row.get("To Date","-"))
                row_index = start_row+i
                if i+2 >= len(table21.rows):  
                    table21.add_row()  # Add rows if needed

                table21.cell(row_index+1, 0).text = str(i+1)  # Serial No.
                table21.cell(row_index+1, 1).text = str(row.get("Project Title", "-"))
                table21.cell(row_index+1, 2).text = str(row.get("Number of Students", "-"))
                table21.cell(row_index+1, 3).text = str(row.get("Title of Hackathon","-"))
                table21.cell(row_index+1, 4).text = str(row.get("Organized By", "-")) 
                table21.cell(row_index+1, 5).text = str(row.get("Date", "-"))  
                table21.cell(row_index+1, 6).text = str(row.get("Status", "-"))
                n=1
                s1_1=1

            row_index=-1
            table21.add_row()
            table21.rows[row_index].cells[0].merge(table21.rows[row_index].cells[-2])
            paragraph = table21.cell(-1, 5).paragraphs[0]
            run = paragraph.add_run("Total score")
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            table21.cell(-1, 6).text = str(n)
            m = m+n
    mentor = m+n

    placeholders = {
            "{{research}}": str(research),
            "{{self}}": str(selfm),
            "{{mentorship}}": str(mentor),
            "{{academics}}": str(academics),  
            "{{name}}":detaillist[0],
            "{{designation}}":detaillist[1],
            "{{dept}}":detaillist[2],
            "{{empid}}":detaillist[3]
        }
    
    placeholders2 = {
            "{{research}}": str(research),
            "{{selfm}}": str(selfm),
            "{{mentor}}": str(mentor),
            "{{academics}}": str(academics),  
        }

    for i in range(1, 14):
        placeholders2[f"{{{{r{i}_1}}}}"] = globals().get(f"r{i}_1", None)

    # Assign pi_1 to placeholders2[pi_1] for i in range 1 to 7
    for i in range(1, 8):
        placeholders2[f"{{{{p{i}_1}}}}"] = globals().get(f"p{i}_1", None)

    # Assign si_1 to placeholders2[si_1] for i in range 1 to 5
    for i in range(1, 6):
        placeholders2[f"{{{{s{i}_1}}}}"] = globals().get(f"s{i}_1", None)
    placeholders2["{{u1}}"]=u1

    score=[academics, research, selfm, mentor,hod]

    fdoc = Document("Faculty Appraisal- Corrective Action Report.docx")

    # Replace placeholders in paragraphs
    for table in fdoc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = paragraph.text
                    for placeholder, value in placeholders2.items():
                        if placeholder in full_text:
                            full_text = full_text.replace(placeholder, str(value))
                    paragraph.text = full_text
    lasttable = fdoc.tables[2]
    assispro = [0.3, 0.3, 0.15, 0.15, 0.1]
    assospro = [0.2, 0.4, 0.15, 0.15, 0.1]
    prof = [0.1, 0.4, 0.2, 0.2, 0.1]
    tot = 0
    print(detaillist[1])

    for row_idx, row in zip(range(2, 6), lasttable.rows[2:6]):  # rows 2 to 5
        for cell_idx, cell in zip(range(1, 6), row.cells[1:6]):  # cells 1 to 5
            if row_idx == 2:
                cell.text = str(score[cell_idx - 1])  

            elif row_idx == 3:
                if detaillist[1] == "Professor":
                    cell.text = str(prof[cell_idx - 1])
                elif detaillist[1] == "Associate Professor":
                    cell.text = str(assospro[cell_idx - 1])
                elif detaillist[1] == "Assistant Professor":
                    cell.text = str(assispro[cell_idx - 1])
                else:
                    cell.text = "0"

            elif row_idx == 4:
                if detaillist[1] == "Professor":
                    weight = prof[cell_idx - 1]
                elif detaillist[1] == "Associate Professor":
                    weight = assospro[cell_idx - 1]
                elif detaillist[1] == "Assistant Professor":
                    weight = assispro[cell_idx - 1]
                else:
                    weight = 0

                weighted_score = score[cell_idx - 1] * weight
                cell.text = str(weighted_score)
                tot += weighted_score

            else:
                print("Not filled")

    # Write total to last cell in 5th row (index 4)
    lasttable.rows[4].cells[-1].text = str(tot)

    

    # Save final doc
    fdoc.save("appfilled_template.docx")
    print("Document saved as filled_template.docx")
        # Replace placeholders in paragraphs with the corresponding values
    for placeholder, value in placeholders.items():
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
    # Save the modified document
    output_doc_path = "filled_template.docx"
    doc.save(output_doc_path)
    fdoc.save("debug_filled_template.docx")
    print(f"Word document saved as {output_doc_path}")

def copy_table_contents(source_table, dest_table):
    """Copy contents from source table to destination table"""
    # Ensure destination table has enough rows
    while len(dest_table.rows) < len(source_table.rows):
        dest_table.add_row()

    # Copy cell contents
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            try:
                dest_table.cell(i, j).text = cell.text
            except IndexError:
                print(f"Warning: Could not copy cell at row {i}, column {j}")

# def process_blueprint(excel_path, staffname):
#     """Process the blueprint and fill the template"""
#     try:
#         # Load both documents
#         source_doc = Document("MSP Self-Appraisal form.docx")
#         template_doc = Document("template.docx")  # Your template document

#         # Copy contents from each table
#         for i, source_table in enumerate(source_doc.tables):
#             try:
#                 # Make sure template has corresponding table
#                 if i < len(template_doc.tables):
#                     copy_table_contents(source_table, template_doc.tables[i])
#             except Exception as e:
#                 print(f"Error copying table {i}: {str(e)}")

        # Fill in placeholders
        placeholders = {
            "{{research}}": str(research),
            "{{self}}": str(selfm),
            "{{mentorship}}": str(mentor),
            "{{name}}": detaillist[0],
            "{{designation}}": detaillist[1],
            "{{dept}}": detaillist[2],
            "{{empid}}": detaillist[3]
        }

        # Replace placeholders in paragraphs
        for paragraph in template_doc.paragraphs:
            for placeholder, value in placeholders.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)

        # Save the filled template
        template_doc.save("filled_template.docx")
        print("Successfully filled template and saved as filled_template.docx")
        
        return True

    # except Exception as e:
    #     print(f"Error processing document: {str(e)}")
    #     return False

if __name__ == '__main__':
    # with app.app_context():
    #     db.create_all()
    app.run(debug=True)